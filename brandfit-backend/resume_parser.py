"""Local resume text and hyperlink extraction for Impulso (PDF / DOCX)."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import urlparse

from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader
from pypdf.errors import FileNotDecryptedError, PdfReadError

MAX_RESUME_BYTES = 5 * 1024 * 1024
MAX_RESUME_CHARS_FOR_AI = 50000
MIN_EXTRACTED_CHARS = 40

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}

PLACEHOLDER_URL_TOKENS = {
    "portfolio",
    "link",
    "website",
    "github",
    "linkedin",
    "url",
    "http",
    "https",
    "www",
    "example",
    "example.com",
}

VISIBLE_URL_RE = re.compile(
    r"(?i)\b(?:https?://|www\.)[^\s<>()\"']+"
)


class ResumeExtractionError(Exception):
    """Raised when a resume file cannot be validated or read."""


@dataclass
class DetectedLink:
    url: str
    anchor_text: str = ""
    source: str = ""
    classification: str = "other"

    def to_dict(self) -> Dict[str, str]:
        return {
            "url": self.url,
            "anchor_text": self.anchor_text,
            "source": self.source,
            "classification": self.classification,
        }


@dataclass
class ExtractedResumeText:
    file_name: str
    text: str
    character_count: int
    truncated_for_ai: bool
    ai_text: str
    warnings: List[str]
    detected_links: List[DetectedLink] = field(default_factory=list)

    @property
    def detected_link_dicts(self) -> List[Dict[str, str]]:
        return [link.to_dict() for link in self.detected_links]


def _extension(filename: str) -> str:
    name = (filename or "").strip().lower()
    if "." not in name:
        return ""
    return "." + name.rsplit(".", 1)[-1]


def validate_resume_upload(
    filename: Optional[str],
    content_type: Optional[str],
    file_bytes: bytes,
) -> str:
    if file_bytes is None or len(file_bytes) == 0:
        raise ResumeExtractionError("The uploaded file is empty.")

    if len(file_bytes) > MAX_RESUME_BYTES:
        raise ResumeExtractionError("Resume must be 5 MB or smaller.")

    ext = _extension(filename or "")
    mime = (content_type or "").split(";")[0].strip().lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise ResumeExtractionError("Unsupported file type. Please upload a PDF or DOCX file.")

    if mime and mime not in ALLOWED_CONTENT_TYPES and mime != "application/octet-stream":
        if ext not in ALLOWED_EXTENSIONS:
            raise ResumeExtractionError("Unsupported file type. Please upload a PDF or DOCX file.")

    return ext


def normalize_url(raw_url: str) -> str:
    value = (raw_url or "").strip()
    if not value:
        return ""

    value = value.strip("<>\"'`")
    value = value.rstrip(".,);:]'")
    value = value.strip()

    if not value:
        return ""

    lowered = value.lower()
    if lowered in PLACEHOLDER_URL_TOKENS:
        return ""

    # Reject bare placeholder-like labels that are not real URLs.
    if "/" not in value and "." not in value:
        return ""

    if lowered.startswith("mailto:"):
        return ""

    if value.startswith("//"):
        value = "https:" + value
    elif re.match(r"(?i)^www\.", value):
        value = "https://" + value
    elif re.match(r"(?i)^[a-z0-9.-]+\.[a-z]{2,}(/.*)?$", value) and not re.match(
        r"(?i)^https?://", value
    ):
        value = "https://" + value

    if not re.match(r"(?i)^https?://", value):
        return ""

    parsed = urlparse(value)
    host = (parsed.netloc or "").lower()
    if not host or host in PLACEHOLDER_URL_TOKENS:
        return ""

    path = (parsed.path or "").strip().lower().strip("/")
    if path in PLACEHOLDER_URL_TOKENS and not parsed.query:
        return ""

    # Drop fragment noise; keep query when present.
    normalized = f"{parsed.scheme}://{parsed.netloc}{parsed.path or ''}"
    if parsed.query:
        normalized += "?" + parsed.query
    return normalized.rstrip("/")


def classify_link(url: str, anchor_text: str = "") -> str:
    parsed = urlparse(url)
    host = (parsed.netloc or "").lower()
    path = (parsed.path or "").strip("/")
    segments = [seg for seg in path.split("/") if seg]
    anchor = (anchor_text or "").strip().lower()

    if "linkedin.com" in host:
        return "linkedin"

    if "github.com" in host:
        if len(segments) >= 2:
            return "project"
        if len(segments) == 1 and segments[0].lower() not in {
            "features",
            "pricing",
            "marketplace",
            "explore",
            "topics",
            "collections",
            "events",
            "sponsors",
            "about",
            "login",
            "join",
        }:
            return "github"
        return "other"

    if host.endswith("github.io") or "gitlab.com" in host or "bitbucket.org" in host:
        if len(segments) >= 1:
            return "project"
        return "portfolio"

    portfolio_suffixes = (
        "vercel.app",
        "netlify.app",
        "pages.dev",
        "herokuapp.com",
        "web.app",
        "firebaseapp.com",
        "carrd.co",
        "notion.site",
        "behance.net",
        "myportfolio.com",
        "cargo.site",
    )
    if any(host == suffix or host.endswith("." + suffix) for suffix in portfolio_suffixes):
        return "portfolio"

    project_keywords = ("repo", "repository", "demo", "project", "github", "live", "app")
    if any(word in anchor for word in project_keywords) and "linkedin" not in host:
        if "github.com" in host or len(segments) >= 1:
            return "project"

    if host and "linkedin.com" not in host and "github.com" not in host:
        # Likely personal/portfolio site when it is a root-ish personal domain.
        if len(segments) <= 1:
            return "portfolio"
        return "project"

    return "other"


def _dedupe_links(links: List[DetectedLink]) -> List[DetectedLink]:
    ordered: List[DetectedLink] = []
    seen = set()
    for link in links:
        key = link.url.lower()
        if not key or key in seen:
            continue
        seen.add(key)
        ordered.append(link)
    return ordered


def _pdf_obj_get(obj: Any, key: str) -> Any:
    try:
        if obj is None:
            return None
        if hasattr(obj, "get"):
            return obj.get(key)
    except Exception:
        return None
    return None


def _pdf_resolve(obj: Any) -> Any:
    try:
        if hasattr(obj, "get_object"):
            return obj.get_object()
    except Exception:
        return obj
    return obj


def _extract_uri_from_annotation(annot_obj: Any) -> str:
    annot_obj = _pdf_resolve(annot_obj)
    if annot_obj is None:
        return ""

    direct_uri = _pdf_obj_get(annot_obj, "/URI")
    if direct_uri:
        return str(direct_uri)

    action = _pdf_obj_get(annot_obj, "/A")
    action = _pdf_resolve(action)
    if action is not None:
        action_uri = _pdf_obj_get(action, "/URI")
        if action_uri:
            return str(action_uri)

    return ""


def _extract_pdf_annotation_links(reader: PdfReader) -> List[DetectedLink]:
    found: List[DetectedLink] = []
    for page in reader.pages:
        annots = _pdf_obj_get(page, "/Annots")
        if not annots:
            continue
        annots = _pdf_resolve(annots)
        try:
            iterable = list(annots)
        except Exception:
            continue

        for annot in iterable:
            try:
                annot_obj = _pdf_resolve(annot)
                uri = _extract_uri_from_annotation(annot_obj)
                normalized = normalize_url(uri)
                if not normalized:
                    continue
                contents = _pdf_obj_get(annot_obj, "/Contents")
                anchor = str(contents) if contents else ""
                found.append(
                    DetectedLink(
                        url=normalized,
                        anchor_text=anchor.strip(),
                        source="pdf_annotation",
                        classification=classify_link(normalized, anchor),
                    )
                )
            except Exception:
                continue
    return found


def _open_pdf_reader(file_bytes: bytes) -> PdfReader:
    try:
        reader = PdfReader(io.BytesIO(file_bytes), strict=False)
    except FileNotDecryptedError as exc:
        raise ResumeExtractionError(
            "This PDF appears to be encrypted or password-protected. Please upload an unlocked file."
        ) from exc
    except PdfReadError as exc:
        raise ResumeExtractionError(
            "The PDF could not be read. It may be corrupted or unreadable."
        ) from exc
    except Exception as exc:
        raise ResumeExtractionError(
            "Failed to open the PDF file. It may be corrupted or unsupported."
        ) from exc

    if getattr(reader, "is_encrypted", False):
        try:
            result = reader.decrypt("")
            if result == 0:
                raise ResumeExtractionError(
                    "This PDF appears to be encrypted or password-protected. Please upload an unlocked file."
                )
        except ResumeExtractionError:
            raise
        except Exception as exc:
            raise ResumeExtractionError(
                "This PDF appears to be encrypted or password-protected. Please upload an unlocked file."
            ) from exc
    return reader


def _extract_pdf_text_and_links(file_bytes: bytes) -> Tuple[str, List[DetectedLink]]:
    reader = _open_pdf_reader(file_bytes)

    chunks: List[str] = []
    try:
        for page in reader.pages:
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""
            if page_text.strip():
                chunks.append(page_text)
    except Exception as exc:
        raise ResumeExtractionError(
            "Failed while extracting text from the PDF. The file may be corrupted."
        ) from exc

    text = "\n".join(chunks).strip()
    if not text:
        raise ResumeExtractionError(
            "No readable text could be extracted. The PDF may be image-only or scanned. "
            "Please upload a text-based resume."
        )

    links = _extract_pdf_annotation_links(reader)
    links.extend(_extract_visible_text_links(text, source="pdf_visible_text"))
    return text, _dedupe_links(links)


def _extract_docx_hyperlinks(document: Document) -> List[DetectedLink]:
    found: List[DetectedLink] = []
    rels = getattr(document.part, "rels", {}) or {}

    # Map rId -> target URL from relationships.
    rel_targets = {}
    for rel in rels.values():
        try:
            reltype = str(getattr(rel, "reltype", "") or "")
            target = str(getattr(rel, "target_ref", "") or "")
            if "hyperlink" in reltype.lower() and target:
                rel_targets[rel.rId] = target
                normalized = normalize_url(target)
                if normalized:
                    found.append(
                        DetectedLink(
                            url=normalized,
                            anchor_text="",
                            source="docx_relationship",
                            classification=classify_link(normalized, ""),
                        )
                    )
        except Exception:
            continue

    def walk_paragraph(paragraph) -> None:
        try:
            element = paragraph._element
        except Exception:
            return
        for hyperlink in element.findall(".//" + qn("w:hyperlink")):
            try:
                r_id = hyperlink.get(qn("r:id"))
                anchor_parts = []
                for node in hyperlink.findall(".//" + qn("w:t")):
                    if node.text:
                        anchor_parts.append(node.text)
                anchor = "".join(anchor_parts).strip()
                target = rel_targets.get(r_id, "") if r_id else ""
                if not target:
                    # External hyperlink sometimes stored differently
                    target = hyperlink.get(qn("w:anchor")) or ""
                normalized = normalize_url(target)
                if not normalized:
                    # Anchor text itself may be a URL label like github.com/user
                    normalized = normalize_url(anchor)
                if not normalized:
                    continue
                found.append(
                    DetectedLink(
                        url=normalized,
                        anchor_text=anchor,
                        source="docx_hyperlink",
                        classification=classify_link(normalized, anchor),
                    )
                )
            except Exception:
                continue

    for paragraph in document.paragraphs:
        walk_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    walk_paragraph(paragraph)

    return found


def _extract_docx_text_and_links(file_bytes: bytes) -> Tuple[str, List[DetectedLink]]:
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ResumeExtractionError(
            "The DOCX file could not be read. It may be corrupted or unreadable."
        ) from exc

    parts: List[str] = []
    for paragraph in document.paragraphs:
        value = (paragraph.text or "").strip()
        if value:
            parts.append(value)

    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts).strip()
    if not text:
        raise ResumeExtractionError(
            "No readable text could be extracted from the DOCX file."
        )

    links = _extract_docx_hyperlinks(document)
    links.extend(_extract_visible_text_links(text, source="docx_visible_text"))
    return text, _dedupe_links(links)


def _extract_visible_text_links(text: str, source: str) -> List[DetectedLink]:
    found: List[DetectedLink] = []
    for match in VISIBLE_URL_RE.finditer(text or ""):
        normalized = normalize_url(match.group(0))
        if not normalized:
            continue
        found.append(
            DetectedLink(
                url=normalized,
                anchor_text="",
                source=source,
                classification=classify_link(normalized, ""),
            )
        )
    return found


# Tokens that must not be altered by spacing repair.
_PROTECTED_TOKENS = [
    "PyTest",
    "MySQL",
    "DevOps",
    "FastAPI",
    "LangChain",
    "GitHub",
    "LinkedIn",
    "CI/CD",
    "N+1",
    "GPT-4",
    "GitLab",
    "OpenAI",
    "TypeScript",
    "JavaScript",
    "Next.js",
    "Node.js",
    "Vue.js",
    "C++",
    "C#",
    ".NET",
    "PyTorch",
    "TensorFlow",
    "PostgreSQL",
    "MongoDB",
    "GraphQL",
    "BrandFit",
]

# High-confidence glued-word repairs only (avoid broad rewriting).
_GLUED_WORD_REPAIRS = (
    (re.compile(r"(?i)\bbackenddriven\b"), "backend-driven"),
    (re.compile(r"(?i)\bfrontenddriven\b"), "frontend-driven"),
    (re.compile(r"(?i)\bdatadriven\b"), "data-driven"),
    (re.compile(r"(?i)\btestdriven\b"), "test-driven"),
    (re.compile(r"(?i)\beventdriven\b"), "event-driven"),
)

_EMAIL_RE = re.compile(r"(?i)\b[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,}\b")
_URL_RE = re.compile(r"(?i)\b(?:https?://|www\.)[^\s<>()\"']+")
_COMPACT_HANDLE_RE = re.compile(r"(?i)(?<![A-Za-z0-9])@[A-Za-z0-9_]{2,}")
_COMPACT_METRIC_RE = re.compile(r"(?i)\b\d+(?:\.\d+)?[KMB]\+?\b")


def _protect_tokens(text: str) -> Tuple[str, Dict[str, str]]:
    protected: Dict[str, str] = {}
    result = text
    for index, token in enumerate(_PROTECTED_TOKENS):
        # Word-boundary aware so we don't swallow suffixes (LangChainfor).
        pattern = re.compile(rf"(?<![A-Za-z0-9]){re.escape(token)}(?![A-Za-z0-9])", re.IGNORECASE)
        if not pattern.search(result):
            continue
        placeholder = f"__IMPULSO_TOKEN_{index}__"
        result = pattern.sub(placeholder, result)
        protected[placeholder] = token
    return result, protected


def _restore_tokens(text: str, protected: Dict[str, str]) -> str:
    result = text
    for placeholder, token in protected.items():
        result = result.replace(placeholder, token)
    return result


def _protect_spans(text: str, pattern: re.Pattern[str], prefix: str) -> Tuple[str, Dict[str, str]]:
    protected: Dict[str, str] = {}
    counter = 0

    def repl(match: re.Match[str]) -> str:
        nonlocal counter
        original = match.group(0)
        placeholder = f"__IMPULSO_{prefix}_{counter}__"
        protected[placeholder] = original
        counter += 1
        return placeholder

    return pattern.sub(repl, text), protected


def repair_email(value: str) -> str:
    """Remove whitespace accidentally inserted inside email addresses."""
    if value is None:
        return ""
    text = str(value).strip()
    if not text:
        return ""
    # Collapse spaces around @ and dots inside a likely email.
    if "@" in text:
        text = re.sub(r"\s*@\s*", "@", text)
        text = re.sub(r"\s*\.\s*", ".", text)
        text = re.sub(r"\s+", "", text) if re.match(r"(?i)^[A-Z0-9._%+\-@.]+$", text) else text
    return text.strip()


def _apply_glued_word_repairs(text: str) -> str:
    result = text
    for pattern, replacement in _GLUED_WORD_REPAIRS:
        result = pattern.sub(replacement, result)

    # Known technical token glued to a following lowercase word: LangChainfor -> LangChain for
    for token in _PROTECTED_TOKENS:
        result = re.sub(
            rf"(?i)\b({re.escape(token)})([a-z]{{2,}})\b",
            r"\1 \2",
            result,
        )
    return result


def normalize_spacing_and_hyphens(text: str) -> str:
    """
    Repair common PDF extraction artifacts without inventing words.
    Never inserts spaces inside URLs, emails, handles, metrics, or known tech tokens.
    """
    if not text:
        return ""

    text = str(text).replace("\r\n", "\n").replace("\r", "\n")
    # Join hyphenated line-break artifacts: "experi-\nence" / "experi- ence"
    text = re.sub(r"([A-Za-z]{2,})-\s*\n\s*([a-z]{2,})", r"\1\2", text)
    text = re.sub(r"([A-Za-z]{2,})-\s{1,}([a-z]{2,})", r"\1\2", text)

    # Fix known glued forms before protecting tokens.
    text = _apply_glued_word_repairs(text)

    # Protect spans that must remain untouched by spacing heuristics.
    text, url_map = _protect_spans(text, _URL_RE, "URL")
    text, email_map = _protect_spans(text, _EMAIL_RE, "EMAIL")
    text, handle_map = _protect_spans(text, _COMPACT_HANDLE_RE, "HANDLE")
    text, metric_map = _protect_spans(text, _COMPACT_METRIC_RE, "METRIC")
    text, token_map = _protect_tokens(text)

    # lowerUpper boundary (does not affect protected tokens/spans).
    text = re.sub(r"([a-z])([A-Z])", r"\1 \2", text)
    # Spaces around percentages, but not inside protected metrics like 10K+.
    text = re.sub(r"([A-Za-z])(\d+%)", r"\1 \2", text)
    text = re.sub(r"(\d+%)([A-Za-z])", r"\1 \2", text)

    text = _restore_tokens(text, token_map)
    text = _restore_tokens(text, metric_map)
    text = _restore_tokens(text, handle_map)
    text = _restore_tokens(text, email_map)
    text = _restore_tokens(text, url_map)
    return text


def clean_extracted_text(text: str) -> str:
    """
    Light cleanup of extracted resume text.
    Does not rewrite meaning or invent content.
    """
    if not text:
        return ""

    # Keep pure URLs/emails/handles/metrics stable if the whole field is one of them.
    stripped = str(text).strip()
    if _URL_RE.fullmatch(stripped) or re.fullmatch(r"(?i)https?://\S+", stripped):
        return normalize_url(stripped) or re.sub(r"\s+", "", stripped)
    if "@" in stripped and stripped.count("@") == 1 and "://" not in stripped:
        repaired = repair_email(stripped)
        if "@" in repaired and "." in repaired.split("@")[-1]:
            return repaired
    if _COMPACT_METRIC_RE.fullmatch(re.sub(r"\s+", "", stripped)):
        return re.sub(r"\s+", "", stripped)
    if _COMPACT_HANDLE_RE.fullmatch(re.sub(r"\s+", "", stripped)):
        return re.sub(r"\s+", "", stripped)

    text = normalize_spacing_and_hyphens(text)
    cleaned_lines: List[str] = []
    for raw_line in text.split("\n"):
        line = raw_line.replace("\t", " ")
        # Remove spaces before punctuation.
        line = re.sub(r"\s+([,.;:!?)\]}])", r"\1", line)
        # Normalize repeated spaces; keep hyphenated terms intact.
        line = re.sub(r" {2,}", " ", line)
        line = line.strip()
        if line:
            cleaned_lines.append(line)

    return "\n".join(cleaned_lines).strip()


def finalize_parsed_profile(
    profile_draft: Dict[str, Any],
    detected_links: Optional[List[Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    """
    Deterministic post-AI validation pass.
    Detected hyperlinks always overwrite AI-provided profile/project URLs.
    """
    if not isinstance(profile_draft, dict):
        return profile_draft

    draft = profile_draft
    links = draft.get("links") if isinstance(draft.get("links"), dict) else {}
    draft["links"] = {
        "linkedin": links.get("linkedin", "") or "",
        "github": links.get("github", "") or "",
        "portfolio": links.get("portfolio", "") or "",
    }

    personal = draft.get("personal") if isinstance(draft.get("personal"), dict) else {}
    if personal:
        email = repair_email(str(personal.get("email", "") or ""))
        personal["email"] = email
        draft["personal"] = personal

    detected = [link for link in (detected_links or []) if isinstance(link, dict)]
    by_class: Dict[str, List[str]] = {
        "linkedin": [],
        "github": [],
        "portfolio": [],
        "project": [],
    }
    for link in detected:
        classification = str(link.get("classification") or "").strip().lower()
        url = normalize_url(str(link.get("url") or ""))
        if not url:
            continue
        if classification in by_class and url not in by_class[classification]:
            by_class[classification].append(url)

    # Always prefer exact detected profile URLs when present.
    if by_class["linkedin"]:
        draft["links"]["linkedin"] = by_class["linkedin"][0]
    else:
        draft["links"]["linkedin"] = normalize_url(str(draft["links"].get("linkedin") or ""))
    if by_class["github"]:
        draft["links"]["github"] = by_class["github"][0]
    else:
        draft["links"]["github"] = normalize_url(str(draft["links"].get("github") or ""))
    if by_class["portfolio"]:
        draft["links"]["portfolio"] = by_class["portfolio"][0]
    else:
        draft["links"]["portfolio"] = normalize_url(str(draft["links"].get("portfolio") or ""))

    projects = draft.get("projects") if isinstance(draft.get("projects"), list) else []
    project_urls = list(by_class["project"])
    used_urls = set()

    def _norm_name(value: str) -> str:
        return re.sub(r"[^a-z0-9]+", "", (value or "").lower())

    for project in projects:
        if not isinstance(project, dict):
            continue
        current_url = normalize_url(str(project.get("url") or ""))
        project_name = str(project.get("name") or "")
        matched = ""

        # Prefer a detected project URL that matches this project name/anchor/repo.
        name_key = _norm_name(project_name)
        for link in detected:
            if str(link.get("classification") or "").lower() != "project":
                continue
            url = normalize_url(str(link.get("url") or ""))
            if not url or url in used_urls:
                continue
            anchor_key = _norm_name(str(link.get("anchor_text") or ""))
            repo_key = _norm_name(url.rstrip("/").split("/")[-1].replace("-", " "))
            if name_key and (
                name_key == anchor_key
                or name_key == repo_key
                or (name_key in _norm_name(url))
                or (anchor_key and anchor_key in name_key)
                or (repo_key and repo_key in name_key)
            ):
                matched = url
                break

        if not matched and current_url and current_url in project_urls:
            matched = current_url

        if matched:
            project["url"] = matched
            used_urls.add(matched)
        else:
            # Do not keep AI-modified URLs that are not in detected project links.
            project["url"] = ""

    # Assign remaining detected project URLs in order to projects still missing URLs.
    remaining = [url for url in project_urls if url not in used_urls]
    for project in projects:
        if not isinstance(project, dict) or project.get("url"):
            continue
        if not remaining:
            break
        project["url"] = remaining.pop(0)

    draft["projects"] = projects
    return draft


_PROJECT_SECTION_HEADERS = re.compile(
    r"(?i)^(projects?|personal projects|selected projects|key projects|side projects)\s*:?\s*$"
)
_SECTION_STOP_HEADERS = re.compile(
    r"(?i)^("
    r"experience|work experience|professional experience|employment|education|"
    r"skills|technical skills|certifications?|awards|publications|summary|"
    r"profile|about|contact|activities|volunteer|interests"
    r")\s*:?\s*$"
)
_FRAGMENT_NAME_START = re.compile(
    r"(?i)^(and|with|using|to|for|that|which|the|a|an|of|in|on|from|by|as|into|over|under)\b"
)


def _looks_like_project_title(line: str) -> bool:
    value = clean_extracted_text(line or "")
    if not value:
        return False
    if len(value) < 3 or len(value) > 140:
        return False
    if value.lower().startswith(("technologies", "tech stack", "built with", "tools")):
        return False
    if re.match(r"^[\d\.\)\-]+$", value):
        return False
    if _FRAGMENT_NAME_START.match(value) and value[:1].islower():
        return False
    if value[:1].islower() and not value.startswith(("iOS", "e-")):
        return False
    if re.search(r"(?i)\b(and|with|or|the|a|an|to|for|of)$", value):
        return False
    return True


def extract_project_blocks(resume_text: str) -> List[Dict[str, str]]:
    """
    Heuristically extract ordered project blocks from a Projects section.
    Used to recover missing project names/descriptions without inventing content.
    """
    if not resume_text:
        return []

    lines = [line.strip() for line in clean_extracted_text(resume_text).splitlines() if line.strip()]
    if not lines:
        return []

    start_index = -1
    for index, line in enumerate(lines):
        if _PROJECT_SECTION_HEADERS.match(line):
            start_index = index + 1
            break

    if start_index < 0:
        return []

    section_lines: List[str] = []
    for line in lines[start_index:]:
        if _SECTION_STOP_HEADERS.match(line):
            break
        section_lines.append(line)

    if not section_lines:
        return []

    blocks: List[Dict[str, str]] = []
    current_name = ""
    current_body: List[str] = []

    def flush() -> None:
        nonlocal current_name, current_body
        body = clean_extracted_text("\n".join(current_body))
        name = clean_extracted_text(current_name)
        if name and _looks_like_project_title(name):
            blocks.append(
                {
                    "name": name,
                    "description": body,
                    "nearby_text": clean_extracted_text(
                        "\n".join([name] + current_body).strip()
                    ),
                }
            )
        current_name = ""
        current_body = []

    bullet_re = re.compile(r"^[-*•·]\s+")
    for line in section_lines:
        stripped = bullet_re.sub("", line).strip()
        is_bullet = bool(bullet_re.match(line))
        looks_like_heading = (not is_bullet) and _looks_like_project_title(stripped)

        if looks_like_heading and (not current_name or current_body):
            if current_name or current_body:
                flush()
            current_name = stripped
            continue

        current_body.append(stripped)

    flush()
    return [block for block in blocks if block.get("name")]


def extract_resume_text(
    filename: Optional[str],
    content_type: Optional[str],
    file_bytes: bytes,
) -> ExtractedResumeText:
    ext = validate_resume_upload(filename, content_type, file_bytes)
    safe_name = (filename or f"resume{ext}").strip() or f"resume{ext}"

    if ext == ".pdf":
        text, links = _extract_pdf_text_and_links(file_bytes)
    else:
        text, links = _extract_docx_text_and_links(file_bytes)

    cleaned = clean_extracted_text(text)

    if len(cleaned) < MIN_EXTRACTED_CHARS:
        raise ResumeExtractionError(
            "Extracted text is too short to parse reliably. "
            "The file may be image-only, sparse, or unreadable."
        )

    # Reclassify after normalization/dedupe for stable labels.
    classified: List[DetectedLink] = []
    for link in links:
        anchor = (link.anchor_text or "").strip()
        # Keep anchors lightweight; avoid aggressive spacing repairs on short labels.
        if anchor and not _URL_RE.fullmatch(anchor):
            anchor = re.sub(r"\s+", " ", anchor).strip()
        classified.append(
            DetectedLink(
                url=normalize_url(link.url) or link.url,
                anchor_text=anchor,
                source=link.source,
                classification=classify_link(link.url, anchor),
            )
        )
    classified = _dedupe_links(classified)

    warnings: List[str] = []
    truncated = len(cleaned) > MAX_RESUME_CHARS_FOR_AI
    ai_text = cleaned[:MAX_RESUME_CHARS_FOR_AI]
    if truncated:
        warnings.append(
            f"Resume text was truncated to {MAX_RESUME_CHARS_FOR_AI} characters before AI parsing."
        )

    return ExtractedResumeText(
        file_name=safe_name,
        text=cleaned,
        character_count=len(cleaned),
        truncated_for_ai=truncated,
        ai_text=ai_text,
        warnings=warnings,
        detected_links=classified,
    )


def prepare_text_for_ai(extracted: ExtractedResumeText) -> Tuple[str, List[str]]:
    return extracted.ai_text, list(extracted.warnings)
